#!/usr/bin/env python3
"""
Pricing Proposal Generator - Matches FUSION template format
"""

import json
import sys
import base64
from datetime import datetime, timedelta
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# MeridianLink Brand Colors
ML_BLUE = RGBColor(0, 73, 142)

def format_currency(amount):
    """Format number as USD currency"""
    return f"${amount:,.2f}"

def shade_cell(cell, color):
    """Add background color to a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_table_with_style(doc, rows, cols):
    """Add a styled table matching the template"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    return table

def main():
    """Main entry point"""
    try:
        # Get proposal data
        if len(sys.argv) < 2:
            raise ValueError("Proposal data required")

        proposal_data = json.loads(sys.argv[1])

        # Create document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run("PRICING PROPOSAL")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = ML_BLUE

        # Prepared For and Date
        prepared = doc.add_paragraph()
        prepared.add_run("Prepared For: ").bold = True
        customer_name = proposal_data.get('customerName', 'N/A')
        contact_name = proposal_data.get('customerContact', '')
        if contact_name:
            prepared.add_run(f"{customer_name} ({contact_name})")
        else:
            prepared.add_run(customer_name)

        date_para = doc.add_paragraph()
        date_para.add_run("Date: ").bold = True
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))

        doc.add_paragraph()  # Spacing

        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        exec_summary = doc.add_paragraph(
            "This proposal outlines the pricing for MeridianLink's Mortgage Loan Origination System (LOS) "
            "integrated with DocMagic for document generation and management. The solution provides a complete "
            "end-to-end origination platform with scalable, per-transaction pricing."
        )
        exec_summary.paragraph_format.space_after = Pt(12)

        # Primary Services Table
        doc.add_heading('Primary Services', level=2)
        line_items = proposal_data.get('lineItems', [])
        platform_fee = proposal_data.get('platformFee', 0)
        contract_years = proposal_data.get('contractTermYears', 1)

        # Add platform fee for Mortgage products - one row per year
        has_mortgage = any('mortgage' in item.get('productName', '').lower() for item in line_items)
        platform_fee_rows = contract_years if (has_mortgage and platform_fee > 0) else 0

        if line_items:
            table = add_table_with_style(doc, rows=len(line_items) + 1 + platform_fee_rows, cols=4)
            table.autofit = False
            table.allow_autofit = False

            # Header row
            header_cells = table.rows[0].cells
            headers = ['Service / Module', 'One-Time Fee', 'Per Transaction', 'Monthly Minimum']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                # Style header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                shade_cell(header_cells[i], "004B8E")  # ML Blue

            # Data rows
            total_setup = 0
            total_monthly = 0

            for idx, item in enumerate(line_items):
                row_cells = table.rows[idx + 1].cells

                product_name = item.get('productName', 'N/A')
                setup_fee = item.get('setupFee', 0)
                per_file_fee = item.get('perFileFee', 0)
                monthly_commitment = item.get('monthlyCommitment', 0)

                total_setup += setup_fee
                total_monthly += monthly_commitment

                row_cells[0].text = product_name
                row_cells[1].text = format_currency(setup_fee)
                row_cells[2].text = format_currency(per_file_fee) if per_file_fee > 0 else "N/A"
                row_cells[3].text = format_currency(monthly_commitment)

                # Center align currency columns
                for col_idx in [1, 2, 3]:
                    for paragraph in row_cells[col_idx].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add platform fee rows for Mortgage products - one row per year
            if has_mortgage and platform_fee > 0:
                for year in range(1, contract_years + 1):
                    platform_row_idx = len(line_items) + year
                    platform_cells = table.rows[platform_row_idx].cells
                    platform_cells[0].text = f"MeridianLink Mortgage Platform Fee - Year {year} (Billed Monthly)"
                    platform_cells[1].text = "N/A"
                    platform_cells[2].text = "N/A"
                    platform_cells[3].text = format_currency(platform_fee)

                    # Center align
                    for col_idx in [1, 2, 3]:
                        for paragraph in platform_cells[col_idx].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                total_monthly += (platform_fee * contract_years)

            # Set column widths
            widths = (Inches(3.5), Inches(1.2), Inches(1.2), Inches(1.2))
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

        doc.add_paragraph()  # Spacing

        # Annual Investment Summary
        doc.add_heading('Annual Investment Summary', level=2)

        # Contract terms and yearly tier info
        contract_years = proposal_data.get('contractTermYears', 1)
        yearly_tiers = proposal_data.get('yearlyTiers', {})

        # Calculate year-by-year investment
        year1_setup = total_setup
        contract_total = 0
        yearly_costs = {}

        for year in range(1, contract_years + 1):
            year_tier = yearly_tiers.get(str(year)) if yearly_tiers else {}
            tier_name = year_tier.get('tierName', 'Standard')

            # Calculate this year's costs - only include line items for this year
            year_monthly = 0
            for item in line_items:
                item_year = item.get('year', 1)
                if item_year == year:
                    year_monthly += item.get('monthlyCommitment', 0)

            # Add platform fee for this year
            year_monthly += platform_fee

            year_total = (year1_setup if year == 1 else 0) + (year_monthly * 12)
            yearly_costs[year] = year_total
            contract_total += year_total

        # Build summary table - one row per year plus header and total
        num_rows = contract_years + 2  # header + years + total
        summary_table = add_table_with_style(doc, rows=num_rows, cols=2)
        summary_table.autofit = False

        # Header
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = "Cost Category"
        header_cells[1].text = "Amount"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            shade_cell(cell, "004B8E")

        # Year-by-year rows
        for year in range(1, contract_years + 1):
            row_idx = year
            row_cells = summary_table.rows[row_idx].cells
            year_tier = yearly_tiers.get(str(year)) if yearly_tiers else {}
            tier_name = year_tier.get('tierName', 'Standard')

            label = f"Year {year} ({tier_name})"
            row_cells[0].text = label
            row_cells[1].text = format_currency(yearly_costs[year])

            # Right align amounts
            for paragraph in row_cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Total row
        total_row_idx = contract_years + 1
        row_cells = summary_table.rows[total_row_idx].cells
        row_cells[0].text = f"Total {contract_years}-Year Investment"
        row_cells[1].text = format_currency(contract_total)

        # Make total row bold
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for paragraph in row_cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths
        widths = (Inches(3.8), Inches(1.8))
        for row in summary_table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        doc.add_paragraph()  # Spacing

        # Contract Terms
        doc.add_heading('Contract Terms', level=2)
        terms = doc.add_paragraph()
        contract_years = proposal_data.get('contractTermYears', 1)
        yearly_tiers = proposal_data.get('yearlyTiers', {})

        terms.add_run("Initial Term: ").bold = True
        terms.add_run(f"{contract_years} year{'s' if contract_years > 1 else ''}\n")

        if yearly_tiers and contract_years > 1:
            terms.add_run("Tier Schedule: ").bold = True
            tier_schedule = []
            for year in range(1, contract_years + 1):
                year_tier = yearly_tiers.get(str(year))
                if year_tier:
                    tier_schedule.append(f"Year {year}: {year_tier.get('tierName', 'Standard')}")
            if tier_schedule:
                terms.add_run(" | ".join(tier_schedule) + "\n")

        terms.add_run("Renewal: ").bold = True
        terms.add_run("Automatic renewal terms will be negotiated as contract end approaches")

        doc.add_paragraph()  # Spacing

        # Confidentiality Notice
        confidential = doc.add_paragraph(
            "This proposal and all materials contained herein are confidential and proprietary to MeridianLink, Inc. "
            "This document is intended solely for the use of the recipient and may not be reproduced, distributed, or "
            "disclosed to third parties without prior written consent. © 2026 MeridianLink, Inc. All rights reserved."
        )
        confidential.paragraph_format.space_before = Pt(12)
        for run in confidential.runs:
            run.font.italic = True
            run.font.size = Pt(9)

        # Save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Encode to base64
        doc_base64 = base64.b64encode(doc_bytes.read()).decode('utf-8')

        # Return success
        result = {
            'success': True,
            'documentData': doc_base64,
            'message': 'Document generated successfully'
        }

        print(json.dumps(result))
        sys.exit(0)

    except Exception as e:
        result = {
            'success': False,
            'error': str(e)
        }
        print(json.dumps(result))
        sys.exit(1)

if __name__ == '__main__':
    main()
